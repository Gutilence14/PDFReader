# from spire.doc import *
# from spire.doc.common import *
from base_output import *
import docx
from docx import Document
import json
import queue
from datetime import datetime
import re
import os
import subprocess





def get_table_text(tbl):  
    text = ""  
    for row in tbl.tr_lst:  # 遍历所有行  
        for tc in row.tc_lst:  # 遍历所有单元格  
            # 提取单元格中的段落  
            for p in tc.p_lst:  
                # 遍历段落中的所有运行  
                for r in p.r_lst:  
                    # 提取运行中的文本  
                    text += r.text
    return text





def analysis_word(word_path, out_dir, use_dict=True):
    # analysisa pdf file

    if not os.path.exists(out_dir):
        os.makedirs(out_dir)

    if word_path.endswith('.doc'):
        subprocess.run(
            ['soffice', '--convert-to', 'docx', '--outdir',
                os.path.dirname(word_path), word_path],
        )
        word_path = word_path[:-4] + ".docx"
        print(word_path)

    if not use_dict:
        res = BaseWordOutput(
            Word=[]
        )
    else:
        res = {
            word_path: {},
        }

    doc = Document(word_path)

    """
    find heading 1 
    """

    # 初始化变量
    current_heading = None
    current_table = None
    sections = {}
    section_content = []

    # 遍历文档中的所有元素
    for elem in doc.element.body:
        if isinstance(elem, docx.oxml.text.paragraph.CT_P):
            # 检查是否为标题1
            if (elem.pPr is not None) and (elem.pPr.pStyle is not None) and (int(elem.pPr.pStyle.val) == 2):
                current_heading = elem.text
                # 如果当前有标题，保存之前的节
                if current_heading:

                    res[word_path][current_heading] = {
                                        "Title": "",
                                        "Text": "",
                                        "Table": "",
                                        "Figure": [],
                    }
                    # sections[current_heading] = {
                    #     'text': section_content, 'tables': current_table}
                    # section_content = []
                    # current_table = None
            else:
                # 如果不是标题，将文本添加到当前节的文本列表中
                if res[word_path].get(current_heading):
                    res[word_path][current_heading]["Text"] += elem.text

        elif isinstance(elem, docx.oxml.table.CT_Tbl):
            # 检查是否为表格
            if res[word_path].get(current_heading):
                table = get_table_text(elem)
                res[word_path][current_heading]["Table"] += table

    # 不要忘记添加文档末尾的最后一节
    if current_heading and section_content:
        sections[current_heading] = {
            'text': section_content, 'tables': current_table}

    # 打印每个标题下的文本和表格
    for heading, content in sections.items():
        print(f"{heading}:")
        for para in content['text']:
            print(f"  - {para}")
        tables = content['tables']
        if tables:
            print("  Tables:")
            for table in tables:
                # 此处可以根据需要添加表格的处理逻辑
                pass

    # text_res = ""
    # for para in doc.paragraphs:
    #     text_res += para.text
    # text_res = re.sub(r'\s+', '', text_res)
    # res[word_path]["Text"] = text_res

    # table_res = ""
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             table_res += cell.text
    # table_res = re.sub(r'\s+', '', table_res)
    # res[word_path]["Table"] = table_res

    # images = []
    # dict_rel = doc.part._rels
    # for rel in dict_rel:
    #     rel = dict_rel[rel]
    #     if "image" in rel.target_ref:
    #         print(rel.target_ref)
    #         filename_with_extension = os.path.basename(word_path)
    #         filename, extension = os.path.splitext(filename_with_extension)
    #         current_time = datetime.now().strftime(
    #             "%Y%m%d%H%M%S%f") + '_' + filename + ".png"
    #         figure_image_name = os.path.join(out_dir, current_time)
    #         images.append(figure_image_name)
    #         with open(figure_image_name, "wb") as f:
    #             f.write(rel.target_part.blob)
    # res[word_path]["Figure"] = images
    return res


if __name__ == '__main__':
    word_path = "BIM/no_heading.docx"
    out_dir = 'img_folder'

    out = analysis_word(word_path, out_dir)
    # print(out)

    # word_path, file_extension = os.path.splitext(word_path)
    # word_path = word_path + '.json'

    # with open(word_path, 'w', encoding='utf-8') as json_file:
    #     json.dump(out, json_file, ensure_ascii=False)
